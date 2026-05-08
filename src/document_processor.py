"""
文档处理模块
负责加载、分割和预处理知识库文档
"""
import os
import logging
from typing import List, Optional
from pathlib import Path

from langchain_community.document_loaders import (
    TextLoader,
    PyPDFLoader,
)
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from docx import Document as DocxDocument

from src.config import config

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """文档处理器，用于加载、分割和预处理知识库文档"""

    def __init__(self):
        """
        初始化文档处理器

        创建文本分割器实例，使用配置中的分块大小、重叠长度等参数
        """
        # 创建递归字符文本分割器实例，用于将长文本分割成指定大小的块
        # chunk_size: 每个文本块的最大长度
        # chunk_overlap: 相邻文本块之间的重叠长度，用于保持上下文连续性
        # length_function: 计算文本长度的函数，这里使用内置的len函数
        # separators: 文本分割符列表，按优先级排序，用于识别分割点
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=config.text_processing.chunk_size,
            chunk_overlap=config.text_processing.chunk_overlap,
            length_function=len,
            separators=config.text_processing.separators,
        )
        self.resolved_resume_file = self._resolve_resume_file()
        if self.resolved_resume_file:
            logger.info(f"简历文件识别完成: {self.resolved_resume_file}")
        else:
            logger.warning("未识别到简历文件，将使用关键词规则进行兜底识别")


    def load_documents(self, file_path: Optional[str] = None) -> List[Document]:
        """
        加载文档

        Args:
            file_path: 文件路径，如果为None则加载整个目录

        Returns:
            文档列表
        """
        if file_path is None:
            # 加载整个目录
            return self._load_directory()
        else:
            # 加载单个文件
            return self._load_single_file(file_path)

    def _load_single_file(self, file_path: str) -> List[Document]:
        """
        加载单个文件

        根据文件扩展名选择合适的加载器加载文档内容

        Args:
            file_path: 要加载的文件路径

        Returns:
            加载得到的文档列表

        Raises:
            ValueError: 当文件格式不被支持时抛出异常
            Exception: 加载过程中出现的其他异常
        """
        file_ext = Path(file_path).suffix.lower()

        try:
            documents = self._load_by_extension(file_path)
            self._annotate_documents(documents, Path(file_path))
            logger.info(f"成功加载文件: {file_path}, 包含 {len(documents)} 个文档")
            return documents

        except Exception as e:
            logger.error(f"加载文件失败 {file_path}: {e}")
            raise

    def _load_directory(self) -> List[Document]:
        """
        加载目录中的所有文档

        首先尝试加载主知识库文件，然后递归加载目录下所有支持的文档文件

        Returns:
            所有加载到的文档列表
        """
        data_dir = Path(config.knowledge_base.data_dir)
        if not data_dir.exists():
            logger.error(f"数据目录不存在: {data_dir}")
            return []

        all_documents: List[Document] = []
        supported = {ext.lower() for ext in config.knowledge_base.supported_extensions}

        for file_path in data_dir.rglob("*"):
            if not file_path.is_file():
                continue
            if file_path.suffix.lower() not in supported:
                continue
            try:
                documents = self._load_single_file(str(file_path))
                all_documents.extend(documents)
            except Exception as e:
                logger.error(f"加载文件失败 {file_path}: {e}")

        logger.info(f"目录文档加载完成，共 {len(all_documents)} 个文档对象")
        return all_documents

    def split_documents(self, documents: List[Document]) -> List[Document]:
        """
        分割文档

        使用文本分割器将长文档切分为较小的块，并为每个块添加元数据信息

        Args:
            documents: 原始文档列表

        Returns:
            分割后的文档列表
        """
        if not documents:
            logger.warning("文档列表为空，无需分割")
            return []

        splits = self.text_splitter.split_documents(documents)
        logger.info(f"文档分割完成: {len(documents)} -> {len(splits)} 个块")

        # 为每个块添加元数据
        for i, split in enumerate(splits):
            split.metadata["chunk_id"] = i
            split.metadata["source"] = split.metadata.get("source", "unknown")

        return splits

    def process_knowledge_base(self) -> List[Document]:
        """
        处理整个知识库

        按顺序执行文档加载和文档分割操作，完成知识库的预处理流程

        Returns:
            处理后的文档列表
        """
        # 加载文档
        raw_documents = self.load_documents()

        if not raw_documents:
            logger.warning("未找到任何文档")
            return []

        # 分割文档
        processed_documents = self.split_documents(raw_documents)

        return processed_documents

    @staticmethod
    def _load_by_extension(file_path: str) -> List[Document]:
        """按扩展名加载文档，避免额外依赖导致加载失败"""
        file_ext = Path(file_path).suffix.lower()
        if file_ext == ".txt":
            return TextLoader(file_path, encoding="utf-8").load()
        if file_ext == ".md":
            return TextLoader(file_path, encoding="utf-8").load()
        if file_ext == ".pdf":
            return PyPDFLoader(file_path).load()
        if file_ext == ".docx":
            doc = DocxDocument(file_path)
            text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            return [Document(page_content=text, metadata={"source": Path(file_path).name})]
        raise ValueError(f"不支持的文件格式: {file_ext}")

    def _resolve_resume_file(self) -> Optional[str]:
        """解析简历文件名：配置优先，不存在则自动发现并回退"""
        data_dir = Path(config.knowledge_base.data_dir)
        if not data_dir.exists():
            return None

        configured_name = (config.knowledge_base.resume_file or "").strip()
        if configured_name:
            configured_path = data_dir / configured_name
            if configured_path.exists() and configured_path.is_file():
                return configured_name
            logger.warning(f"RESUME_FILE 指定文件不存在: {configured_path}，将尝试自动发现")

        supported = {ext.lower() for ext in config.knowledge_base.supported_extensions}
        candidates: List[Path] = []
        for file_path in data_dir.rglob("*"):
            if file_path.is_file() and file_path.suffix.lower() in supported:
                candidates.append(file_path)

        if not candidates:
            return None

        keyword_candidates = [p for p in candidates if "简历" in p.stem or "resume" in p.stem.lower() or "cv" in p.stem.lower()]
        target_pool = keyword_candidates if keyword_candidates else candidates
        latest = max(target_pool, key=lambda p: p.stat().st_mtime)
        return latest.name

    def _annotate_documents(self, documents: List[Document], file_path: Path):
        """为文档打标签，供后续检索优先级使用"""
        source_path = str(file_path)
        source_name = file_path.name
        source_type = "other"
        source_priority = 99

        if (self.resolved_resume_file and source_name == self.resolved_resume_file) or "简历" in file_path.stem:
            source_type = "resume"
            source_priority = 1
        elif "additional_docs" in str(file_path.parent):
            source_type = "additional"
            source_priority = 2

        for doc in documents:
            doc.metadata["source"] = source_name
            doc.metadata["source_path"] = source_path
            doc.metadata["source_type"] = source_type
            doc.metadata["source_priority"] = source_priority
